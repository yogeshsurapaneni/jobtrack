import io
import re
import markdown
import weasyprint
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Preprocessor for list formatting
# ---------------------------------------------------------------------------

def preprocess_markdown_lists(text):
    """
    Normalizes markdown text by splitting inline continuous bullet points (e.g. * A * B)
    into line-by-line formatting, and ensures there is a blank line before any list
    block starts so standard markdown parsers render it as a <ul>/<li> block.
    """
    if not text:
        return ""
    
    # 1. Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # 2. Split lines containing multiple bullet points inline
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            processed_lines.append(line)
            continue
            
        # Match ' * ' or ' - ' or ' • ' (surrounded by space) but avoid splitting '**'
        parts = re.split(r'\s+[\-\*\u2022]\s+', line)
        if len(parts) > 1:
            starts_with_bullet = line.lstrip().startswith('* ') or line.lstrip().startswith('- ') or line.lstrip().startswith('• ')
            for i, part in enumerate(parts):
                part_stripped = part.strip()
                if not part_stripped:
                    continue
                if i == 0 and not starts_with_bullet:
                    processed_lines.append(part)
                else:
                    # Clean any trailing/leading asterisks from split points
                    processed_lines.append(f"- {part_stripped}")
        else:
            # Also handle if bullet markers are like "*Point A *Point B" (no space before bullet)
            # but avoid breaking bold markers. Match " *[a-zA-Z]" or " -[a-zA-Z]"
            sub_parts = re.split(r'\s+[\-\*\u2022](?=[a-zA-Z0-9])', line)
            if len(sub_parts) > 1:
                starts_with_bullet = line.lstrip().startswith('*') or line.lstrip().startswith('-') or line.lstrip().startswith('•')
                for i, part in enumerate(sub_parts):
                    part_stripped = part.strip()
                    if not part_stripped:
                        continue
                    if i == 0 and not starts_with_bullet:
                        processed_lines.append(part)
                    else:
                        processed_lines.append(f"- {part_stripped}")
            else:
                processed_lines.append(line)
                
    # 3. Ensure list block starts are preceded by a blank line
    final_lines = []
    for i, line in enumerate(processed_lines):
        stripped = line.strip()
        is_bullet = stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('• ') or re.match(r'^\d+\.\s+', stripped)
        if is_bullet and i > 0:
            prev_stripped = final_lines[-1].strip()
            prev_is_bullet = prev_stripped.startswith('- ') or prev_stripped.startswith('* ') or prev_stripped.startswith('• ') or re.match(r'^\d+\.\s+', prev_stripped)
            if prev_stripped and not prev_is_bullet:
                final_lines.append('')
        final_lines.append(line)
        
    return '\n'.join(final_lines)


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------

def markdown_to_pdf(md_text):
    """
    Converts markdown text to PDF bytes using WeasyPrint.
    Renders a polished, ATS-ready, print-quality document.
    """
    # Normalize bullet points and line-by-line format first
    md_text = preprocess_markdown_lists(md_text)
    
    html_body = markdown.markdown(
        md_text,
        extensions=['tables', 'nl2br']
    )

    full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<style>
  /* ── Page setup ── */
  @page {{
    size: letter;
    margin: 0.65in 0.7in 0.65in 0.7in;
    @bottom-center {{
      content: counter(page) " of " counter(pages);
      font-family: 'Arial', sans-serif;
      font-size: 7.5pt;
      color: #9ca3af;
    }}
  }}

  /* ── Base typography ── */
  body {{
    font-family: 'Arial', 'Helvetica Neue', Helvetica, sans-serif;
    font-size: 10.5pt;
    line-height: 1.5;
    color: #1f2937;
    margin: 0;
    padding: 0;
  }}

  /* ── Name / H1 ── */
  h1 {{
    text-align: center;
    font-size: 22pt;
    font-weight: 700;
    color: #111827;
    margin: 0 0 4px 0;
    padding: 0;
    letter-spacing: 0.8px;
    text-transform: uppercase;
  }}

  /* ── Section headers / H2 ── */
  h2 {{
    font-size: 10pt;
    font-weight: 700;
    color: #1d4ed8;
    text-transform: uppercase;
    letter-spacing: 1px;
    margin: 18px 0 4px 0;
    padding-bottom: 3px;
    border-bottom: 1.5px solid #1d4ed8;
  }}

  /* ── Job title / H3 ── */
  h3 {{
    font-size: 10.5pt;
    font-weight: 700;
    color: #111827;
    margin: 10px 0 2px 0;
  }}

  /* ── H4 — date/location metadata ── */
  h4 {{
    font-size: 9.5pt;
    font-weight: 400;
    color: #6b7280;
    margin: 0 0 4px 0;
  }}

  /* ── Body paragraphs ── */
  p {{
    margin: 0 0 5px 0;
    line-height: 1.5;
  }}

  /* ── Bullet lists — properly indented ── */
  ul {{
    margin: 2px 0 6px 0;
    padding-left: 0;
    list-style: none;
  }}
  ul li {{
    position: relative;
    padding-left: 14px;
    margin-bottom: 3px;
    line-height: 1.45;
    text-align: left;
  }}
  ul li::before {{
    content: "\\2022";   /* bullet */
    position: absolute;
    left: 0;
    top: 0;
    color: #1d4ed8;
    font-weight: 700;
  }}

  /* Nested lists */
  ul ul {{
    margin-top: 2px;
    padding-left: 14px;
  }}

  /* ── Horizontal rule ── */
  hr {{
    border: none;
    border-top: 1px solid #e5e7eb;
    margin: 10px 0;
  }}

  /* ── Inline formatting ── */
  strong {{ font-weight: 700; color: #111827; }}
  em {{ font-style: italic; color: #374151; }}

  /* ── Contact line (centered, small) ── */
  p:first-of-type,
  .contact {{
    text-align: center;
    font-size: 9pt;
    color: #4b5563;
    margin-bottom: 10px;
  }}

  /* ── Tables (for skill grids, if any) ── */
  table {{
    width: 100%;
    border-collapse: collapse;
    margin-bottom: 8px;
    font-size: 10pt;
  }}
  td, th {{
    padding: 3px 6px;
    vertical-align: top;
  }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    return weasyprint.HTML(string=full_html).write_pdf()


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def markdown_to_docx(md_text):
    """
    Converts markdown to a professionally styled DOCX file.
    Bullet alignment is correct via proper hanging-indent paragraph format.
    """
    md_text = preprocess_markdown_lists(md_text)
    doc = docx.Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    # Base Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)   # #1f2937

    # Ensure List Bullet style exists and is tuned
    try:
        bullet_style = doc.styles['List Bullet']
    except KeyError:
        bullet_style = doc.styles.add_style('List Bullet', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Arial'
    bullet_style.font.size = Pt(10.5)

    lines = md_text.split('\n')

    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            _add_empty_para(doc)
            continue

        if stripped.startswith('# '):
            _add_h1(doc, stripped[2:])
        elif stripped.startswith('## '):
            _add_h2(doc, stripped[3:])
        elif stripped.startswith('### '):
            _add_h3(doc, stripped[4:])
        elif stripped.startswith('#### '):
            _add_h4(doc, stripped[5:])
        elif stripped.startswith('- ') or stripped.startswith('* '):
            _add_bullet(doc, stripped[2:])
        elif stripped.startswith('---') or stripped.startswith('***'):
            _add_hr(doc)
        else:
            _add_body(doc, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _add_empty_para(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)


def _add_h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.name  = 'Arial'
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(17, 24, 39)


def _add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text.upper())
    run.font.name  = 'Arial'
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(29, 78, 216)   # #1d4ed8

    # Add bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')        # 0.5pt
    bottom.set(qn('w:space'), '3')
    bottom.set(qn('w:color'), '1d4ed8')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name  = 'Arial'
    run.font.size  = Pt(10.5)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(17, 24, 39)


def _add_h4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(107, 114, 128)


def _add_bullet(doc, text):
    """Bullet with proper hanging indent."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent        = Inches(0.25)
    p.paragraph_format.first_line_indent  = Inches(-0.25)
    _parse_inline(p, text)


def _add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _parse_inline(p, text)


def _add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'e5e7eb')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_inline(paragraph, text):
    """Parse **bold**, *italic*, and plain text inline."""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**') and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from PDF, DOCX, or TXT file bytes.
    """
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'txt':
        return file_bytes.decode('utf-8', errors='ignore')
    elif ext == 'docx':
        import io
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs)
    elif ext == 'pdf':
        try:
            import io
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
            return "\n".join(text_parts)
        except ImportError:
            raise RuntimeError("PDF parsing requires 'pypdf' package to be installed.")
    else:
        raise ValueError(f"Unsupported file extension: .{ext}")
